"""
Complete document processing pipeline.
Combines OCR, chunking, embedding, and indexing.
"""

import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import hashlib
import json

# Use absolute imports
import sys
from pathlib import Path

# Add src to path for imports
src_dir = Path(__file__).parent.parent.parent
sys.path.insert(0, str(src_dir))

try:
    from core.ocr import OCRProcessor, OCRFactory
    from core.chunking import ChunkingFactory, Chunk
    from core.embedding import EmbeddingFactory
    from core.reranking import RerankingFactory
    from models.document import Document, DocumentMetadata, DocumentSource, DocumentType
except ImportError:
    # Fallback for direct execution
    from ..ocr import OCRProcessor, OCRFactory
    from ..chunking import ChunkingFactory, Chunk
    from ..embedding import EmbeddingFactory
    from ..reranking import RerankingFactory
    from ...models.document import Document, DocumentMetadata, DocumentSource, DocumentType

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Complete document processing pipeline."""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize document processor.
        
        Args:
            config: Processing configuration
        """
        self.config = config or self._default_config()
        
        # Initialize components
        self.ocr_processor = None
        self.chunker = None
        self.embedder = None
        self.reranker = None
        
        self._initialize_components()
    
    def _default_config(self) -> Dict[str, Any]:
        """Get default configuration."""
        return {
            "ocr": {
                "enabled": True,
                "engine": "tesseract",
                "language": "eng"
            },
            "chunking": {
                "strategy": "recursive",
                "max_chunk_size": 512,
                "overlap": 50
            },
            "embedding": {
                "model": "minilm",
                "cache_enabled": True,
                "batch_size": 32
            },
            "reranking": {
                "enabled": True,
                "model": "cross-encoder",
                "top_k": 10
            },
            "text_cleaning": {
                "enabled": True,
                "remove_excess_whitespace": True,
                "normalize_unicode": True,
                "min_text_length": 50
            }
        }
    
    def _initialize_components(self):
        """Initialize processing components."""
        # OCR
        if self.config["ocr"]["enabled"]:
            try:
                self.ocr_processor = OCRFactory.create_processor(
                    engine=self.config["ocr"]["engine"],
                    language=self.config["ocr"]["language"]
                )
                logger.info("OCR processor initialized")
            except Exception as e:
                logger.warning(f"Failed to initialize OCR: {e}")
                self.ocr_processor = None
        
        # Chunking
        chunking_config = self.config["chunking"]
        self.chunker = ChunkingFactory.create_strategy(
            strategy=chunking_config["strategy"],
            max_chunk_size=chunking_config["max_chunk_size"]
        )
        logger.info(f"Chunker initialized: {chunking_config['strategy']}")
        
        # Embedding
        embedding_config = self.config["embedding"]
        try:
            self.embedder = EmbeddingFactory.create_model(
                model_type=embedding_config["model"]
            )
            logger.info(f"Embedder initialized: {embedding_config['model']}")
        except Exception as e:
            logger.warning(f"Failed to initialize embedder: {e}")
            self.embedder = None
        
        # Reranking
        if self.config["reranking"]["enabled"]:
            try:
                self.reranker = RerankingFactory.create_reranker(
                    reranker_type=self.config["reranking"]["model"]
                )
                logger.info(f"Reranker initialized: {self.config['reranking']['model']}")
            except Exception as e:
                logger.warning(f"Failed to initialize reranker: {e}")
                self.reranker = None
    
    def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Tuple[Document, List[Chunk]]:
        """
        Process a file through the complete pipeline.
        
        Args:
            file_path: Path to the file
            metadata: Optional document metadata
            
        Returns:
            Tuple of (Document, list of Chunks with embeddings)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Processing file: {file_path}")
        
        # Step 1: Extract text (with OCR if needed)
        text, extraction_metadata = self._extract_text(file_path)
        
        if not text or len(text.strip()) < self.config["text_cleaning"]["min_text_length"]:
            raise ValueError(f"File contains too little text: {len(text or '')} characters")
        
        # Step 2: Clean text
        cleaned_text = self._clean_text(text)
        
        # Step 3: Create document
        document = self._create_document(
            file_path=file_path,
            content=cleaned_text,
            extraction_metadata=extraction_metadata,
            user_metadata=metadata
        )
        
        # Step 4: Chunk document
        chunks = self.chunker.chunk(cleaned_text)
        logger.info(f"Created {len(chunks)} chunks")
        
        # Step 5: Add embeddings to chunks
        if self.embedder:
            chunks = self._add_embeddings_to_chunks(chunks)
        
        return document, chunks
    
    def _extract_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from file, using OCR if needed."""
        # Check file type
        suffix = file_path.suffix.lower()
        
        # Text-based files
        if suffix in ['.txt', '.md', '.json', '.csv', '.xml', '.html']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                return text, {
                    "method": "direct_read",
                    "success": True,
                    "file_type": suffix[1:]
                }
            except UnicodeDecodeError:
                # Try different encoding
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        text = f.read()
                    return text, {
                        "method": "direct_read_latin1",
                        "success": True,
                        "file_type": suffix[1:],
                        "encoding": "latin-1"
                    }
                except Exception as e:
                    return "", {"method": "direct_read", "success": False, "error": str(e)}
        
        # PDF files
        elif suffix == '.pdf':
            if self.ocr_processor:
                return self.ocr_processor.extract_text(str(file_path))
            else:
                # Try direct PDF extraction
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(str(file_path))
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                    
                    if text.strip():
                        return text, {
                            "method": "pdf_direct",
                            "success": True,
                            "file_type": "pdf"
                        }
                except ImportError:
                    pass
                
                return "", {"method": "pdf", "success": False, "error": "OCR not available"}
        
        # Office documents
        elif suffix in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return text, {
                    "method": "docx",
                    "success": True,
                    "file_type": "docx"
                }
            except ImportError:
                return "", {"method": "docx", "success": False, "error": "python-docx not installed"}
            except Exception as e:
                return "", {"method": "docx", "success": False, "error": str(e)}
        
        elif suffix in ['.xlsx', '.xls']:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                text_parts = []
                
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    for row in ws.iter_rows(values_only=True):
                        row_text = " ".join([str(cell) for cell in row if cell])
                        if row_text:
                            text_parts.append(row_text)
                
                text = "\n".join(text_parts)
                return text, {
                    "method": "xlsx",
                    "success": True,
                    "file_type": "xlsx"
                }
            except ImportError:
                return "", {"method": "xlsx", "success": False, "error": "openpyxl not installed"}
            except Exception as e:
                return "", {"method": "xlsx", "success": False, "error": str(e)}
        
        # Image files (require OCR)
        elif suffix in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            if self.ocr_processor:
                return self.ocr_processor.extract_text(str(file_path))
            else:
                return "", {"method": "image", "success": False, "error": "OCR not available"}
        
        # Unsupported format
        else:
            return "", {
                "method": "unknown",
                "success": False,
                "error": f"Unsupported file format: {suffix}"
            }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        if not self.config["text_cleaning"]["enabled"]:
            return text
        
        cleaned = text
        
        # Remove excess whitespace
        if self.config["text_cleaning"]["remove_excess_whitespace"]:
            import re
            # Replace multiple spaces with single space
            cleaned = re.sub(r'\s+', ' ', cleaned)
            # Remove leading/trailing whitespace from each line
            cleaned = '\n'.join([line.strip() for line in cleaned.split('\n')])
        
        # Normalize unicode (if needed)
        if self.config["text_cleaning"]["normalize_unicode"]:
            try:
                import unicodedata
                cleaned = unicodedata.normalize('NFKC', cleaned)
            except ImportError:
                pass
        
        return cleaned.strip()
    
    def _create_document(
        self,
        file_path: Path,
        content: str,
        extraction_metadata: Dict[str, Any],
        user_metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """Create a Document object from processed content."""
        from uuid import uuid4
        from datetime import datetime
        
        # Generate document ID
        doc_id = str(uuid4())
        
        # Calculate content hash
        content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
        
        # Determine document source and type
        source, doc_type = self._infer_source_and_type(file_path, extraction_metadata)
        
        # Create metadata
        metadata_dict = {
            "source": source.value,
            "document_type": doc_type.value,
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "extraction_method": extraction_metadata.get("method", "unknown"),
            "extraction_success": extraction_metadata.get("success", False),
            "processing_timestamp": datetime.utcnow().isoformat()
        }
        
        # Add user metadata if provided
        if user_metadata:
            metadata_dict.update(user_metadata)
        
        # Add extraction metadata details
        for key in ["language", "engine", "pages", "characters_extracted"]:
            if key in extraction_metadata:
                metadata_dict[f"extraction_{key}"] = extraction_metadata[key]
        
        metadata = DocumentMetadata(**metadata_dict)
        
        # Create document title from filename
        title = file_path.stem.replace('_', ' ').replace('-', ' ')
        
        return Document(
            id=doc_id,
            title=title,
            content=content,
            content_hash=content_hash,
            metadata=metadata,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
    
    def _infer_source_and_type(self, file_path: Path, extraction_metadata: Dict[str, Any]) -> Tuple[DocumentSource, DocumentType]:
        """Infer document source and type from file and metadata."""
        # Default values
        source = DocumentSource.LOCAL_FILE
        doc_type = DocumentType.TEXT
        
        # Infer from file extension
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            doc_type = DocumentType.PDF
        elif suffix in ['.docx', '.doc']:
            doc_type = DocumentType.DOCX
        elif suffix in ['.xlsx', '.xls']:
            doc_type = DocumentType.XLSX
        elif suffix in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            doc_type = DocumentType.IMAGE
        
        # Infer from extraction metadata
        extraction_method = extraction_metadata.get("method", "")
        if "gmail" in extraction_method:
            source = DocumentSource.GMAIL
            doc_type = DocumentType.EMAIL
        elif "chat" in extraction_method:
            source = DocumentSource.GOOGLE_CHAT
            doc_type = DocumentType.CHAT_MESSAGE
        elif "drive" in extraction_method:
            source = DocumentSource.GOOGLE_DRIVE
        
        return source, doc_type
    
    def _add_embeddings_to_chunks(self, chunks: List[Chunk]) -> List[Chunk]:
        """Add embeddings to document chunks."""
        if not self.embedder:
            return chunks
        
        # Extract chunk texts
        chunk_texts = [chunk.text for chunk in chunks]
        
        # Generate embeddings
        batch_size = self.config["embedding"]["batch_size"]
        embeddings = self.embedder.embed(chunk_texts, batch_size=batch_size)
        
        # Add embeddings to chunks
        for chunk, embedding in zip(chunks, embeddings):
            if embedding:
                chunk.metadata["embedding"] = embedding
                chunk.metadata["embedding_dim"] = len(embedding)
                chunk.metadata["embedding_model"] = self.embedder.model_name
        
        logger.info(f"Added embeddings to {len([c for c in chunks if 'embedding' in c.metadata])} chunks")
        return chunks
    
    def rerank_search_results(self, query: str, documents: List[str], top_k: Optional[int] = None) -> List[Tuple[int, float]]:
        """
        Rerank search results.
        
        Args:
            query: Search query
            documents: List of document texts
            top_k: Number of top results to return (None for all)
            
        Returns:
            List of (index, score) pairs
        """
        if not self.reranker or not documents:
            # Fallback: return documents in original order with dummy scores
            return [(i, 1.0 - (i * 0.01)) for i in range(len(documents))]
        
        # Rerank documents
        results = self.reranker.rerank(query, documents)
        
        # Limit to top_k if specified
        if top_k is not None:
            results = results[:top_k]
        
        return results


class ProcessingPipeline:
    """Orchestrates the complete document processing pipeline."""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.processor = DocumentProcessor(config)
        self.stats = {
            "documents_processed": 0,
            "chunks_created": 0,
            "embeddings_generated": 0,
            "errors": 0
        }
    
    def process(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Optional[Tuple[Document, List[Chunk]]]:
        """
        Process a document through the complete pipeline.
        
        Args:
            file_path: Path to the file
            metadata: Optional document metadata
            
        Returns:
            Tuple of (Document, Chunks) or None if processing failed
        """
        try:
            document, chunks = self.processor.process_file(file_path, metadata)
            
            # Update statistics
            self.stats["documents_processed"] += 1
            self.stats["chunks_created"] += len(chunks)
            self.stats["embeddings_generated"] += len([c for c in chunks if "embedding" in c.metadata])
            
            logger.info(f"Processed document: {document.title} ({len(chunks)} chunks)")
            return document, chunks
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            self.stats["errors"] += 1
            return None
    
    def get_stats(self) -> Dict[str, Any]:
        """Get processing statistics."""
        return self.stats.copy()


# Test the processing pipeline
if __name__ == "__main__":
    import sys
    
    print("Testing Document Processing Pipeline")
    print("=" * 60)
    
    try:
        # Create pipeline
        pipeline = ProcessingPipeline()
        
        # Test with a sample text file
        test_file = Path(__file__).parent / "test_sample.txt"
        
        # Create test file if it doesn't exist
        if not test_file.exists():
            test_content = """
            SULV Construction Project Report
            ================================
            
            Project: St Ives Duplex
            Client: Yvette
            Project ID: SULV-2024-001
            
            Overview:
            SULV Construction is building a two-story duplex in St Ives.
            The project includes sustainable features like solar panels
            and rainwater harvesting.
            
            Timeline:
            - Site preparation: 2 weeks
            - Foundation: 3 weeks  
            - Framing: 4 weeks
            - Finishing: 8 weeks
            
            Budget: $850,000
            Status: Planning phase
            """
            
            test_file.write_text(test_content)
            print(f"Created test file: {test_file}")
        
        # Process the file
        print(f"\nProcessing test file: {test_file}")
        result = pipeline.process(str(test_file), {"project_id": "SULV-2024-001"})
        
        if result:
            document, chunks = result
            print(f"✅ Document processed: {document.title}")
            print(f"  Chunks: {len(chunks)}")
            print(f"  With embeddings: {sum(1 for c in chunks if 'embedding' in c.metadata)}")
        else:
            print("❌ Processing failed")
        
        print(f"\nPipeline stats: {pipeline.get_stats()}")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()